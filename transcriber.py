import gc
import os
import io
import subprocess
import threading
import time
import wave
import torch
import whisperx
from docx import Document
from dotenv import load_dotenv

load_dotenv()

# Global model cache — one entry per model name, loaded on first use
_models: dict = {}
_model_lock = threading.Lock()

_device = "cuda" if torch.cuda.is_available() else "cpu"
_compute_type = "float16" if _device == "cuda" else "int8"

MODEL_FAST    = "base"
MODEL_PRECISE = "large-v3"
MODEL_MISTRAL = "voxtral-mini-2507"


def _get_model(model_name: str):
    with _model_lock:
        if model_name not in _models:
            _models[model_name] = whisperx.load_model(
                model_name, _device, compute_type=_compute_type
            )
    return _models[model_name]


def _unload_models():
    with _model_lock:
        _models.clear()
    gc.collect()
    if torch.cuda.is_available():
        torch.cuda.empty_cache()


def convert_to_wav(input_path, output_path):
    """Convert audio file to 16 kHz mono WAV using ffmpeg."""
    result = subprocess.run(
        [
            "ffmpeg", "-y",
            "-i", input_path,
            "-ar", "16000",
            "-ac", "1",
            output_path,
        ],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(f"ffmpeg failed: {result.stderr}")


def segments_to_text(segments):
    """Convert diarized segments to speaker-labeled plain text."""
    lines = []
    current_speaker = None
    for segment in segments:
        speaker = segment.get("speaker", "UNKNOWN")
        text = segment["text"].strip()
        if speaker != current_speaker:
            if current_speaker is not None:
                lines.append("")
            lines.append(f"[{speaker}]")
            current_speaker = speaker
        lines.append(text)
    return "\n".join(lines)


def segments_to_docx(segments, filename):
    """Convert diarized segments to DOCX bytes with bold speaker headings."""
    doc = Document()
    doc.add_heading(filename, level=1)
    current_speaker = None
    current_para = None
    for segment in segments:
        speaker = segment.get("speaker", "UNKNOWN")
        text = segment["text"].strip()
        if speaker != current_speaker:
            para = doc.add_paragraph()
            run = para.add_run(f"[{speaker}]")
            run.bold = True
            current_speaker = speaker
            current_para = doc.add_paragraph(text)
        else:
            current_para.add_run(" " + text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def text_to_docx(text, filename):
    """Convert plain text to DOCX bytes (no speaker labels)."""
    doc = Document()
    doc.add_heading(filename, level=1)
    for paragraph in text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def run_transcription(job, wav_path, filename, model_name: str = MODEL_PRECISE):
    """Dispatch to the appropriate transcription backend."""
    if model_name == MODEL_MISTRAL:
        _run_mistral(job, wav_path, filename)
    else:
        _run_whisperx(job, wav_path, filename, model_name)


def _run_whisperx(job, wav_path, filename, model_name: str):
    """Run full WhisperX pipeline and update job dict in-place."""
    hf_token = os.environ.get("HF_TOKEN")
    if not hf_token:
        job["status"] = "error"
        job["error"] = "HF_TOKEN not set"
        return

    def cancelled():
        return job.get("cancelled", False)

    try:
        job["progress"] = "Loading model..."
        model = _get_model(model_name)

        if cancelled():
            job["status"] = "cancelled"
            return

        job["progress"] = "Transcribing audio..."
        audio = whisperx.load_audio(wav_path)
        audio_duration_s = len(audio) / 16000.0

        t0 = time.time()
        result = model.transcribe(audio, batch_size=4)
        transcription_s = time.time() - t0

        language = result["language"]
        job["language"] = language

        if cancelled():
            job["status"] = "cancelled"
            return

        job["progress"] = f"Aligning timestamps..."
        align_model, metadata = whisperx.load_align_model(
            language_code=language, device=_device
        )
        result = whisperx.align(
            result["segments"], align_model, metadata, audio, _device,
            return_char_alignments=False,
        )

        if cancelled():
            job["status"] = "cancelled"
            return

        job["progress"] = "Diarizing speakers..."
        diarize_model = whisperx.diarize.DiarizationPipeline(
            token=hf_token, device=_device
        )
        t0 = time.time()
        diarize_segments = diarize_model(audio)
        diarization_s = time.time() - t0
        result = whisperx.assign_word_speakers(diarize_segments, result)

        if cancelled():
            job["status"] = "cancelled"
            return

        job["progress"] = "Generating output files..."
        segments = result["segments"]
        job["txt"] = segments_to_text(segments)
        job["docx"] = segments_to_docx(segments, filename)
        job["stats"] = {
            "audio_duration_s": audio_duration_s,
            "transcription_s": transcription_s,
            "diarization_s": diarization_s,
        }
        job["status"] = "done"
        job["progress"] = "Complete"

    except Exception as exc:
        job["status"] = "error"
        job["error"] = str(exc)
    finally:
        # Unload all models and free VRAM
        _unload_models()
        # Clean up temp WAV
        try:
            os.remove(wav_path)
        except OSError:
            pass


def _run_mistral(job, wav_path, filename):
    """Run Mistral Voxtral API transcription and update job dict in-place."""
    api_key = os.environ.get("MISTRAL_API_KEY")
    if not api_key:
        job["status"] = "error"
        job["error"] = "MISTRAL_API_KEY not set"
        return

    def cancelled():
        return job.get("cancelled", False)

    try:
        from mistralai import Mistral

        # Measure audio duration from WAV header without loading full audio
        with wave.open(wav_path, "rb") as wf:
            audio_duration_s = wf.getnframes() / float(wf.getframerate())

        if cancelled():
            job["status"] = "cancelled"
            return

        job["progress"] = "Transcribing with Mistral Voxtral..."
        client = Mistral(api_key=api_key)
        t0 = time.time()
        with open(wav_path, "rb") as f:
            response = client.audio.transcriptions.create(
                model=MODEL_MISTRAL,
                file=("audio.wav", f, "audio/wav"),
            )
        transcription_s = time.time() - t0

        if cancelled():
            job["status"] = "cancelled"
            return

        language = getattr(response, "language", None)
        if language:
            job["language"] = language

        job["progress"] = "Generating output files..."
        text = response.text
        job["txt"] = text
        job["docx"] = text_to_docx(text, filename)
        job["stats"] = {
            "audio_duration_s": audio_duration_s,
            "transcription_s": transcription_s,
            "diarization_s": 0,
        }
        job["status"] = "done"
        job["progress"] = "Complete"

    except Exception as exc:
        job["status"] = "error"
        job["error"] = str(exc)
    finally:
        try:
            os.remove(wav_path)
        except OSError:
            pass
